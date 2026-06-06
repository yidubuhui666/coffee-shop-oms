"""Generate the English course-design report (.docx).

Conforms to the spec in the requirements file:
  - title: 14pt Times New Roman, bold
  - body : 10pt Times New Roman
  - 7 sections: Division of labor / Overview / DB design + SQL /
                Instructions for use / Screenshots / Summary / References
"""
import os
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT  = r"C:\Users\许yq\coffee-shop-oms"
DOCS  = os.path.join(ROOT, "docs")
SHOTS = os.path.join(DOCS, "screenshots")
ER    = os.path.join(DOCS, "ER_Diagram.png")
OUT   = r"D:\cxdownload\092Xuyueqian-Coffee Shop Order Management System.docx"

doc = Document()

# ---------- default style: Times New Roman 10pt ----------
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)

def H(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def P(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p

def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    return p

def IMG(path, width=5.8, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(caption)
            r.italic = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(9)
    else:
        P(f"[Image not found: {path}]")

# ---------------- cover ----------------
def cover():
    for _ in range(2): doc.add_paragraph()
    for line, sz, bold in [
        ("Southwest University", 16, True),
        ("School of Computer & Information Science", 12, False),
        ("Database Principles and Applications", 12, False),
        ("Course Design Report", 12, False),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.bold = bold
        r.font.name = "Times New Roman"; r.font.size = Pt(sz)
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Coffee Shop Order Management System")
    r.bold = True; r.font.size = Pt(22); r.font.name = "Times New Roman"
    for _ in range(3): doc.add_paragraph()
    for line in [
        "Team Members:",
        "    114  Yang Fan       (Team Leader)",
        "    092  Xu Yueqian",
        "",
        "Major  : Computer Science and Technology, Class of 2024",
        "Course : Database Principles and Applications",
        f"Date   : {date.today().strftime('%B %Y')}",
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.name = "Times New Roman"; r.font.size = Pt(11)
    doc.add_page_break()

cover()

# ---------------- 1. Division of Labor ----------------
H("1. Division of Labor")
P("1.1 Team Composition.", bold=True)
P("This course design is jointly completed by a two-person team. Both members are "
  "sophomores of the Class of 2024 majoring in Computer Science and Technology at "
  "Southwest University. Both members participated in the entire stack — "
  "requirement discussion, database design, back-end and front-end development, "
  "and final integration testing. The division below reflects task ownership "
  "rather than strict boundaries: Yang Fan (Student ID ending in 114) serves as "
  "team leader with a back-end / database focus, while Xu Yueqian (Student ID "
  "ending in 092) focuses more on the front-end and application packaging side. "
  "An important characteristic of this project is that each member maintains an "
  "independent local MySQL database; code is synchronised via GitHub, and any "
  "database change is paired with an idempotent seed / migration script so the "
  "collaborator can re-sync with a single git pull plus running the script.")

P("1.2 Responsibility Matrix.", bold=True)
tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Phase / Workload"
hdr[1].text = "Yang Fan (114, Leader)"
hdr[2].text = "Xu Yueqian (092)"
for left, a, b in [
    ("Requirement Analysis",                   "Lead",         "Co-author"),
    ("ER Modeling",                            "Lead",         "Review & revise"),
    ("Schema (DDL), Indexes, Views",           "Lead",         "Review"),
    ("Triggers / Stored Procedures",           "Lead",         "Test cases"),
    ("Demo Data Bootstrap Script",             "Lead",         "Co-author"),
    ("Back-end: Orders / Inventory / Stats",   "Lead",         "Co-author"),
    ("Back-end: Menu / Customers / Promos",    "Co-author",    "Lead"),
    ("Front-end: Dashboard / Order pages",     "Lead",         "Co-author"),
    ("Front-end: Menu / Customer / Recipe",    "Co-author",    "Lead"),
    ("CSS Styling & Layout (main theme)",      "Review",       "Lead"),
    ("Product Images (AI-generated)",          "Review",       "Lead"),
    ("DB Migration / Seed Scripts",            "Co-author",    "Lead"),
    ("One-click Launchers (.bat / .command)",  "Co-author",    "Lead"),
    ("Independent local MySQL maintenance",    "Own DB",       "Own DB"),
    ("Integration Testing",                    "Co-author",    "Co-author"),
    ("Report Writing & Formatting",            "Co-author",    "Lead"),
    ("Interface Screenshots",                  "Review",       "Lead"),
]:
    row = tbl.add_row().cells
    row[0].text = left; row[1].text = a; row[2].text = b
for row in tbl.rows:
    for cell in row.cells:
        for par in cell.paragraphs:
            for r in par.runs:
                r.font.name = "Times New Roman"; r.font.size = Pt(10)

P("1.3 Collaboration Workflow.", bold=True)
P("The team adopted Git + GitHub for version control. A repository was created on "
  "GitHub; both members cloned it locally and synchronised changes via the standard "
  "pull / commit / push cycle. A notable choice of this project is that each "
  "member maintains an independent local MySQL database instance — only the source "
  "code is shared through GitHub; the data itself is never committed (the .gitignore "
  "excludes *.db files, the virtual environment and the MySQL credential file "
  ".env.bat). To keep two independent databases consistent, every schema or data "
  "change is shipped together with an idempotent Python seed / migration script "
  "(seed_*.py or migration_*.py). The collaborator only needs to run git pull plus "
  "execute the script once, after which both databases hold the same state. "
  "Commits follow the Conventional-Commits convention (feat:, fix:, docs:, etc.). "
  "A project-level CLAUDE.md records this hard rule as a checklist for every push.")

# ---------------- 2. Overview ----------------
H("2. Overview")
P("2.1 Project Background.", bold=True)
P("The Coffee Shop Order Management System (OMS) is a small but complete business "
  "information system aimed at digitising the daily operations of an independent "
  "coffee shop. The scenario is inspired by a campus coffee shop near Southwest "
  "University, whose owner currently relies on paper order slips and spreadsheet "
  "bookkeeping, both of which are error-prone and provide no analytical "
  "capability.")

P("2.2 Main tasks tackled by this project:", bold=True)
for line in [
    "- Building a relational schema that fully describes the data of a coffee shop "
    "(customers, staff, menu, orders, inventory, promotions, recipes, member "
    "discounts).",
    "- Implementing referential integrity, business rules, indexes, views, "
    "triggers and a stored procedure inside the database itself, so that the data "
    "layer remains consistent regardless of which client manipulates it.",
    "- Providing a browser-based application with role-based access for "
    "Admin / Barista / Cashier that supports CRUD on every entity and exposes "
    "rich multi-table JOIN reports for daily revenue and best-seller statistics.",
    "- Demonstrating how the application interacts with the database through "
    "an ORM (SQLAlchemy), while the underlying SQL still leverages classical DBMS "
    "features.",
]:
    p = doc.add_paragraph(line, style="List Bullet")
    for r in p.runs: r.font.name = "Times New Roman"; r.font.size = Pt(10)

P("2.3 Key problems to solve:", bold=True)
for line in [
    "modelling one-to-many and many-to-many relations without redundancy "
    "and with proper normalisation (BCNF / 3NF);",
    "keeping order totals consistent when order items change (solved by a trigger);",
    "awarding member loyalty points automatically when an order is paid;",
    "preventing over-selling by tracking ingredient inventory through the "
    "many-to-many product-ingredient recipe and warning the operator below "
    "alert thresholds;",
    "supporting member-level discount and promotion-code discount, with "
    "the more favourable of the two applied at checkout;",
    "allowing customers to redeem loyalty points as cash deduction at the "
    "next order (closing the loyalty loop).",
]:
    p = doc.add_paragraph("- " + line)
    for r in p.runs: r.font.name = "Times New Roman"; r.font.size = Pt(10)

# ---------------- 3. Database design ----------------
H("3. Database Design and SQL Statements")
P("3.1 Requirement Analysis.", bold=True)
P("Through interviews with the owner of a campus coffee shop, the following data "
  "flows were identified: a cashier creates an order on behalf of a (possibly "
  "anonymous walk-in) customer; the order consists of one or more items, each "
  "linked to a menu product; the order transitions through "
  "PENDING - PAID - PREPARING - READY - COMPLETED, or it can be CANCELLED; once "
  "paid, the customer earns one loyalty point per RMB spent; each product consumes "
  "ingredients tracked in the inventory; promotions may apply discount codes; the "
  "member level (BRONZE / SILVER / GOLD / PLATINUM) also implies a discount rate.")

P("3.2 Conceptual Design (ER Model).", bold=True)
P("Nine entities are identified: Customer, Staff, Category, Product, Order, "
  "OrderItem, Inventory, Promotion, MemberDiscount, and the associative entity "
  "ProductIngredient. The ER diagram is shown below (Figure 1).")
IMG(ER, width=5.8, caption="Figure 1.  ER Diagram of the Coffee Shop OMS")

P("3.3 Logical Design (Relations).", bold=True)
P("The conceptual model is mapped to ten relations. Primary keys are marked PK; "
  "foreign keys are marked FK.")
for line in [
    "1.  customers       (customer_id PK, name, phone, email, member_level, points, registered_at)",
    "2.  staff           (staff_id PK, username, password, name, role, phone, hire_date, active)",
    "3.  categories      (category_id PK, name, sort_order)",
    "4.  products        (product_id PK, category_id FK, name, description, price, available, stock, image)",
    "5.  promotions      (promo_id PK, code, name, discount, start_date, end_date, active)",
    "6.  orders          (order_id PK, customer_id FK, staff_id FK, promo_id FK, order_time, "
    "                     total_amount, paid_amount, status, payment_method, dine_in, remark, points_used)",
    "7.  order_items     (item_id PK, order_id FK, product_id FK, quantity, unit_price, size, customization)",
    "8.  inventory       (ingredient_id PK, name, unit, quantity, alert_threshold, updated_at)",
    "9.  product_ingredient (product_id PK/FK, ingredient_id PK/FK, consume_qty)",
    "10. member_discount (level PK, discount, label)",
]:
    CODE(line)

P("3.4 Physical Design (MySQL DDL).", bold=True)
P("The full DDL is provided in sql/schema.sql. Representative excerpt of the "
  "orders table:")
CODE("CREATE TABLE orders (")
CODE("    order_id        INT AUTO_INCREMENT PRIMARY KEY,")
CODE("    customer_id     INT,")
CODE("    staff_id        INT NOT NULL,")
CODE("    promo_id        INT,")
CODE("    order_time      DATETIME DEFAULT CURRENT_TIMESTAMP,")
CODE("    total_amount    DECIMAL(10,2) NOT NULL DEFAULT 0,")
CODE("    paid_amount     DECIMAL(10,2) NOT NULL DEFAULT 0,")
CODE("    status          ENUM('PENDING','PAID','PREPARING','READY','COMPLETED','CANCELLED'),")
CODE("    payment_method  ENUM('CASH','CARD','WECHAT','ALIPAY') DEFAULT 'WECHAT',")
CODE("    dine_in         TINYINT(1) DEFAULT 1,")
CODE("    points_used     INT DEFAULT 0,")
CODE("    FOREIGN KEY (customer_id) REFERENCES customers(customer_id) ON DELETE SET NULL,")
CODE("    FOREIGN KEY (staff_id)    REFERENCES staff(staff_id),")
CODE("    FOREIGN KEY (promo_id)    REFERENCES promotions(promo_id) ON DELETE SET NULL,")
CODE("    INDEX idx_time(order_time), INDEX idx_status(status)")
CODE(") ENGINE=InnoDB;")

P("3.5 Constraints and business rules.", bold=True)
for line in [
    "Primary-key and foreign-key constraints guarantee entity and referential integrity.",
    "CHECK (quantity > 0) on order_items prevents zero / negative line items.",
    "ENUM types restrict status, role, payment_method to valid values.",
    "Unique indexes on customers.phone, staff.username, promotions.code prevent duplicates.",
    "Index on orders(order_time) and orders(status) accelerates the dashboard and filter queries.",
    "ON DELETE SET NULL on customer_id / promo_id of orders keeps historical orders valid when a member or promotion is deleted.",
]:
    p = doc.add_paragraph("- " + line)
    for r in p.runs: r.font.name = "Times New Roman"; r.font.size = Pt(10)

P("3.6 Database Programming (Triggers + Stored Procedure).", bold=True)
P("Three triggers and one stored procedure implement business rules inside the "
  "database:")
CODE("-- Trigger 1: recompute order total whenever an item is inserted")
CODE("CREATE TRIGGER trg_item_after_insert AFTER INSERT ON order_items")
CODE("FOR EACH ROW")
CODE("BEGIN")
CODE("    UPDATE orders SET total_amount = (")
CODE("        SELECT COALESCE(SUM(quantity * unit_price),0)")
CODE("        FROM   order_items WHERE order_id = NEW.order_id")
CODE("    ) WHERE order_id = NEW.order_id;")
CODE("END;")
CODE("")
CODE("-- Trigger 2: award loyalty points when an order becomes PAID")
CODE("CREATE TRIGGER trg_order_after_update AFTER UPDATE ON orders")
CODE("FOR EACH ROW")
CODE("BEGIN")
CODE("    IF NEW.status = 'PAID' AND OLD.status <> 'PAID' AND NEW.customer_id IS NOT NULL THEN")
CODE("        UPDATE customers SET points = points + FLOOR(NEW.total_amount)")
CODE("        WHERE customer_id = NEW.customer_id;")
CODE("    END IF;")
CODE("END;")
CODE("")
CODE("-- Trigger 3: deduct ingredient stock when an order item is inserted")
CODE("CREATE TRIGGER trg_item_deduct_stock AFTER INSERT ON order_items")
CODE("FOR EACH ROW")
CODE("BEGIN")
CODE("    UPDATE inventory inv")
CODE("    JOIN   product_ingredient pi ON pi.ingredient_id = inv.ingredient_id")
CODE("    SET    inv.quantity = inv.quantity - pi.consume_qty * NEW.quantity")
CODE("    WHERE  pi.product_id = NEW.product_id;")
CODE("END;")
CODE("")
CODE("-- Stored procedure: place a single-item order in one call")
CODE("CREATE PROCEDURE sp_place_order(IN p_customer INT, IN p_staff INT,")
CODE("                                IN p_product INT,  IN p_qty INT,")
CODE("                                OUT p_order_id INT)")
CODE("BEGIN")
CODE("    DECLARE v_price DECIMAL(8,2);")
CODE("    SELECT price INTO v_price FROM products WHERE product_id = p_product;")
CODE("    INSERT INTO orders(customer_id, staff_id) VALUES (p_customer, p_staff);")
CODE("    SET p_order_id = LAST_INSERT_ID();")
CODE("    INSERT INTO order_items(order_id, product_id, quantity, unit_price)")
CODE("        VALUES (p_order_id, p_product, p_qty, v_price);")
CODE("    UPDATE orders SET status = 'PAID' WHERE order_id = p_order_id;")
CODE("END;")

P("Two views are also defined to simplify reporting:")
CODE("CREATE VIEW v_daily_revenue AS")
CODE("    SELECT DATE(order_time) d, COUNT(*) cnt, SUM(total_amount) revenue")
CODE("    FROM   orders WHERE status <> 'CANCELLED'")
CODE("    GROUP  BY DATE(order_time);")
CODE("")
CODE("CREATE VIEW v_low_stock AS")
CODE("    SELECT name, unit, quantity, alert_threshold")
CODE("    FROM   inventory WHERE quantity <= alert_threshold;")

P("3.7 Representative SQL queries used by the application.", bold=True)
CODE("-- (a) today's revenue")
CODE("SELECT SUM(total_amount) FROM orders")
CODE("WHERE  DATE(order_time) = CURDATE() AND status <> 'CANCELLED';")
CODE("")
CODE("-- (b) top-5 best-selling products (JOIN products with order_items)")
CODE("SELECT p.name, SUM(oi.quantity) AS qty")
CODE("FROM   order_items oi JOIN products p ON p.product_id = oi.product_id")
CODE("GROUP  BY p.product_id ORDER BY qty DESC LIMIT 5;")
CODE("")
CODE("-- (c) member spend ranking (3-table JOIN)")
CODE("SELECT c.name, c.member_level,")
CODE("       SUM(oi.quantity * oi.unit_price) AS spent")
CODE("FROM   customers c")
CODE("JOIN   orders o     ON o.customer_id = c.customer_id")
CODE("JOIN   order_items oi ON oi.order_id   = o.order_id")
CODE("WHERE  o.status <> 'CANCELLED'")
CODE("GROUP  BY c.customer_id ORDER BY spent DESC LIMIT 10;")
CODE("")
CODE("-- (d) ingredient consumption forecast (4-table outer-join)")
CODE("SELECT inv.name, SUM(pi.consume_qty * oi.quantity) AS used_7d")
CODE("FROM   inventory inv")
CODE("LEFT JOIN product_ingredient pi ON pi.ingredient_id = inv.ingredient_id")
CODE("LEFT JOIN order_items oi        ON oi.product_id    = pi.product_id")
CODE("LEFT JOIN orders o              ON o.order_id       = oi.order_id")
CODE("WHERE  o.order_time >= DATE_SUB(CURDATE(), INTERVAL 7 DAY)")
CODE("       AND o.status <> 'CANCELLED'")
CODE("GROUP  BY inv.ingredient_id;")

P("3.8 Sample DML used during testing:", bold=True)
CODE("INSERT INTO orders(customer_id, staff_id) VALUES (1, 3);")
CODE("SET @o := LAST_INSERT_ID();")
CODE("INSERT INTO order_items(order_id, product_id, quantity, unit_price)")
CODE("VALUES (@o, 3, 2, 25.00), (@o, 11, 1, 32.00);")
CODE("UPDATE orders SET status = 'PAID' WHERE order_id = @o;")
CODE("-- trg_item_after_insert  recomputes total_amount to 82.00")
CODE("-- trg_order_after_update awards 82 loyalty points to customer 1")
CODE("-- trg_item_deduct_stock  reduces matching ingredients in inventory")

# ---------------- 4. Instructions ----------------
H("4. Brief Instructions for Use")
P("4.1 Environment Prerequisites: ", bold=True)
P("Python 3.10+, MySQL 8.0+. Dependencies are listed in src/requirements.txt and "
  "include Flask, Flask-SQLAlchemy and PyMySQL.")

P("4.2 One-click Launch.", bold=True)
P("The submission ships with two double-clickable launchers. On Windows, double-"
  "click  启动.bat ; on macOS, double-click  启动咖啡店系统.command . The script "
  "checks Python, creates a local virtual environment on first run, installs the "
  "dependencies and starts the Flask development server on port 5050. A browser "
  "tab automatically opens at http://127.0.0.1:5050/.")

P("4.3 Manual Launch (MySQL mode):", bold=True)
CODE("mysql -uroot -p < sql/schema.sql")
CODE("mysql -uroot -p coffee_shop < sql/data.sql")
CODE("set DATABASE_URL=mysql+pymysql://root:<pwd>@127.0.0.1:3306/coffee_shop?charset=utf8mb4")
CODE("cd src && python app.py    # http://127.0.0.1:5050")

P("4.4 Demo accounts (pre-loaded by the bootstrap routine):", bold=True)
tbl2 = doc.add_table(rows=1, cols=4); tbl2.style = "Light Grid Accent 1"
hdr2 = tbl2.rows[0].cells
hdr2[0].text = "Role"; hdr2[1].text = "Username"; hdr2[2].text = "Password"; hdr2[3].text = "Name"
for r, u, pwd, n in [
    ("ADMIN",   "admin",   "admin123",   "Yang Fan"),
    ("ADMIN",   "xyq",     "070203",     "Xu Yueqian"),
    ("BARISTA", "barista", "barista123", "Li Ming"),
    ("CASHIER", "cashier", "cashier123", "Wang Lei"),
]:
    row = tbl2.add_row().cells
    row[0].text = r; row[1].text = u; row[2].text = pwd; row[3].text = n
for row in tbl2.rows:
    for cell in row.cells:
        for par in cell.paragraphs:
            for r in par.runs:
                r.font.name = "Times New Roman"; r.font.size = Pt(10)

P("4.5 Typical user workflow.", bold=True)
P("Cashier opens 'New Order', picks the customer (or 'Walk-in'), adds menu items "
  "with size and quantity, optionally enters a promotion code or redeems loyalty "
  "points, chooses payment and submits. The Barista watches the 'Orders' page and "
  "updates status to PREPARING / READY / COMPLETED. The Admin manages menu, "
  "inventory, recipes and reviews the Statistics page.")

# ---------------- 5. Screenshots ----------------
H("5. Main Interface Screenshots")
for fname, caption in [
    ("01_dashboard.png",  "Figure 2.  Dashboard - KPI cards and top-5 sellers"),
    ("02_menu.png",       "Figure 3.  Menu - categorised product cards with search & filter"),
    ("03_new_order.png",  "Figure 4.  New Order - multi-line cart with size, payment and discount"),
    ("04_orders.png",     "Figure 5.  Order list - today/history tabs, status filter and pagination"),
    ("05_stats.png",      "Figure 6.  Statistics - daily revenue, product ranking and 8 multi-JOIN reports"),
]:
    IMG(os.path.join(SHOTS, fname), width=5.8, caption=caption)

# ---------------- 6. Summary ----------------
H("6. Summary")
P("6.1 What we accomplished.", bold=True)
P("This course design delivered a fully working Coffee Shop Order Management "
  "System backed by MySQL and a Flask web application. From the database side, "
  "the project exercised the entire design life-cycle: requirement analysis, "
  "conceptual modelling with an ER diagram, logical mapping to 3NF relations, "
  "and physical implementation with DDL/DML, indexes, views, triggers and a "
  "stored procedure. From the application side, the system implements role-based "
  "authentication, complete CRUD on every entity, a multi-line order cart, an "
  "order-status state machine, ingredient inventory with low-stock alerts, "
  "member-level discount, promotion codes, loyalty-point redemption, and rich "
  "daily / per-product statistics with eight multi-table JOIN reports.")

P("6.2 Engineering improvements made during iteration.", bold=True)
for line in [
    "Refactored a 1000-line monolithic app.py into eleven Flask Blueprint modules "
    "(auth, main, products, orders, customers, staff, inventory, promotions, "
    "recipes, analytics, api) under an application factory, improving readability "
    "and team collaboration.",
    "Eliminated N+1 query patterns on the orders and dashboard pages using "
    "joinedload / selectinload, cutting per-page SQL counts from ~70 to ~5.",
    "Introduced server-side pagination for the order and customer lists.",
    "Added 404 / 500 friendly error pages and a smoke-test suite.",
    "All database changes are paired with idempotent seed scripts so that the "
    "collaborator can re-sync with one git pull plus running the script.",
]:
    p = doc.add_paragraph("- " + line)
    for r in p.runs: r.font.name = "Times New Roman"; r.font.size = Pt(10)

P("6.3 Existing problems and future improvements.", bold=True)
for line in [
    "Passwords are currently stored in plain text for grading convenience and "
    "should be replaced with salted hashes (werkzeug.security or bcrypt) before "
    "production use.",
    "The system runs as a single Flask process; for higher concurrency it should "
    "be deployed behind Gunicorn + Nginx.",
    "Front-end styling is intentionally lightweight; a richer SPA (Vue / React) "
    "and chart library (ECharts) could improve UX, especially for the statistics "
    "page on tablets.",
    "An operation-log table could be added to record who edited which record at "
    "what time, further exercising the multi-table relations theme of the course.",
]:
    p = doc.add_paragraph("- " + line)
    for r in p.runs: r.font.name = "Times New Roman"; r.font.size = Pt(10)

P("6.4 Learning outcomes.", bold=True)
P("Through this project the team gained hands-on experience in translating a real "
  "business scenario into a normalised relational model, in using triggers, views "
  "and stored procedures to keep business invariants close to the data, and in "
  "building a complete data-driven web application end to end. The two-person "
  "collaboration also exercised Git-based version control, branch synchronisation, "
  "and code review skills that are essential for future industry work.")

# ---------------- 7. References ----------------
H("7. References")
for ref in [
    "[1] Abraham Silberschatz, Henry F. Korth, S. Sudarshan. Database System Concepts (7th Edition). McGraw-Hill, 2019.",
    "[2] Wang Shan, Sa Shixuan. An Introduction to Database Systems (5th Edition). Higher Education Press, 2014.",
    "[3] Oracle Corporation. MySQL 8.0 Reference Manual. https://dev.mysql.com/doc/refman/8.0/en/",
    "[4] Pallets Projects. Flask Documentation (3.x). https://flask.palletsprojects.com/",
    "[5] Mike Bayer. SQLAlchemy 2.0 Documentation. https://docs.sqlalchemy.org/",
    "[6] Martin Fowler. Patterns of Enterprise Application Architecture. Addison-Wesley, 2002.",
]:
    P(ref)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("WROTE:", OUT)
